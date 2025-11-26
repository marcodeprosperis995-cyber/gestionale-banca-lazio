import streamlit as st
import pandas as pd
from docx import Document
import io
import datetime
import os

# --- CONFIGURAZIONE ---
st.set_page_config(page_title="Gestionale Banca Centro Lazio", page_icon="🏦")
COLORE_AZIENDALE = "#0056b3"

# --- FUNZIONI ---
def carica_dati():
    if os.path.exists("database.xlsx"):
        return pd.read_excel("database.xlsx")
    return None

def crea_word_in_memoria(dati):
    doc = Document()
    # Intestazione
    doc.add_heading('Scheda Personale - Banca Centro Lazio', 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Nome e Cognome: ").bold = True
    p.add_run(f"{dati['Nome']} {dati['Cognome']}\n")
    p.add_run(f"Ruolo: ").bold = True
    p.add_run(f"{dati['Ruolo']}\n")
    p.add_run(f"Email: ").bold = True
    p.add_run(f"{dati['Email']}\n")
    tel = dati['Telefono'] if 'Telefono' in dati else "N/D"
    p.add_run(f"Telefono: ").bold = True
    p.add_run(f"{tel}")
    
    doc.add_paragraph("\n" * 2)
    doc.add_paragraph("Si certifica che i dati sopra riportati sono corretti.")
    
    doc.add_paragraph("\n" * 4)
    paragrafo_firme = doc.add_paragraph()
    data_oggi = datetime.date.today().strftime('%d/%m/%Y')
    paragrafo_firme.add_run(f"Data: {data_oggi}").bold = True
    paragrafo_firme.add_run("\t\t\tFirma Responsabile: __________________")
    
    # Salva in memoria (RAM) invece che su disco
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFACCIA GRAFICA ---

# Intestazione con Logo
col1, col2 = st.columns([1, 4])
with col1:
    if os.path.exists("logo.jpg"):
        st.image("logo.jpg", width=100)
    else:
        st.write("🏦")
with col2:
    st.title("Banca Centro Lazio")
    st.caption("Portale Gestione Risorse Umane")

st.divider()

# Caricamento Excel
df = carica_dati()

if df is None:
    st.error("⚠️ File 'database.xlsx' non trovato! Assicurati di averlo caricato su GitHub.")
else:
    # Barra di ricerca
    ricerca = st.text_input("🔍 Cerca dipendente (Nome o Cognome)", placeholder="Es. Rossi...")

    if ricerca:
        # Filtra dati
        risultati = df[
            df['Nome'].astype(str).str.lower().str.contains(ricerca.lower()) | 
            df['Cognome'].astype(str).str.lower().str.contains(ricerca.lower())
        ]
        
        if risultati.empty:
            st.warning("Nessun risultato trovato.")
        else:
            st.success(f"Trovati {len(risultati)} risultati:")
            
            for index, row in risultati.iterrows():
                with st.container():
                    # Creiamo due colonne per ogni risultato: Testo a sinistra, Bottone a destra
                    c1, c2 = st.columns([3, 1])
                    with c1:
                        st.subheader(f"{row['Nome']} {row['Cognome']}")
                        st.write(f"**Ruolo:** {row['Ruolo']} | **Email:** {row['Email']}")
                    with c2:
                        st.write("") # Spaziatura
                        # Prepariamo il file
                        file_word = crea_word_in_memoria(row)
                        
                        # Tasto Download
                        st.download_button(
                            label="📥 Scarica Word",
                            data=file_word,
                            file_name=f"Scheda_{row['Nome']}_{row['Cognome']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=index
                        )
                    st.divider()
